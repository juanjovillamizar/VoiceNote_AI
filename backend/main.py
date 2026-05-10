import asyncio
import json
import os
import re
import secrets
import subprocess
import tempfile
from contextlib import asynccontextmanager
from datetime import datetime, timedelta, timezone
from io import BytesIO
from typing import AsyncGenerator, Optional

import bcrypt
import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi import (
    Depends,
    FastAPI,
    File,
    HTTPException,
    Query,
    UploadFile,
    BackgroundTasks,
    Request,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from groq import Groq
from jose import JWTError, jwt
from pydantic import BaseModel
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    HRFlowable,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    KeepTogether,
)
from reportlab.platypus.flowables import Flowable
from sqlalchemy import Column, DateTime, ForeignKey, Integer, String, Text, select
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker, create_async_engine
from sqlalchemy.orm import DeclarativeBase, relationship

# ── Imports para Rate Limiting (Seguridad) ─────────────────────────────────────
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded

# ── Config ─────────────────────────────────────────────────────────────────────

DATABASE_URL = os.getenv(
    "DATABASE_URL", "postgresql+asyncpg://voicenote:vnpass@localhost:5432/voicenotedb"
)
JWT_SECRET = os.getenv("JWT_SECRET", "dev-secret-change-this")
CLAUDE_API_KEY = os.getenv("CLAUDE_API_KEY", "")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
ALGORITHM = "HS256"
TOKEN_EXPIRE_DAYS = 7

# Permite localhost para desarrollo. En producción añade tu dominio real al .env
ALLOWED_ORIGINS = os.getenv(
    "ALLOWED_ORIGINS", "http://localhost:3000,http://127.0.0.1:3000"
).split(",")

_groq_client: Optional[Groq] = None


def get_groq() -> Groq:
    global _groq_client
    if _groq_client is None:
        if not GROQ_API_KEY:
            raise HTTPException(
                503, "GROQ_API_KEY no configurada. Agrégala al .env y reinicia."
            )
        _groq_client = Groq(api_key=GROQ_API_KEY)
    return _groq_client


GLOSSARY: dict[str, str] = {
    "doker": "Docker",
    "docker": "Docker",
    "gitt": "Git",
    "kubernets": "Kubernetes",
    "apai": "API",
    "paiton": "Python",
    "yava": "Java",
    "reac": "React",
    "noud": "Node",
    "sequel": "SQL",
    "postgrest": "PostgreSQL",
}

FILLERS = re.compile(
    r"\b(eh+|ah+|mm+|uhh?|este|o sea|o sea que|pues|bueno|digamos|entonces este|osea)\b[,.]?\s*",
    re.IGNORECASE,
)

engine = create_async_engine(DATABASE_URL, echo=False)
AsyncSessionLocal = async_sessionmaker(engine, expire_on_commit=False)
http_bearer = HTTPBearer()

# Inicializar Limitador de Peticiones
limiter = Limiter(key_func=get_remote_address)


# ── ORM ────────────────────────────────────────────────────────────────────────


class Base(DeclarativeBase):
    pass


class User(Base):
    __tablename__ = "users"
    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(100), nullable=False)
    email = Column(String(255), unique=True, nullable=False, index=True)
    password_hash = Column(String(255), nullable=False)
    created_at = Column(DateTime, default=datetime.utcnow)
    meetings = relationship(
        "Meeting", back_populates="user", cascade="all, delete-orphan"
    )


class Meeting(Base):
    __tablename__ = "meetings"
    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), nullable=False)
    title = Column(String(255), nullable=False)
    transcript = Column(Text, default="")
    summary = Column(Text, default="")
    duration = Column(Integer, default=0)
    share_token = Column(String(64), unique=True, nullable=True, index=True)
    created_at = Column(DateTime, default=datetime.utcnow)
    user = relationship("User", back_populates="meetings")


# ── Lifespan y Setup de la App ──────────────────────────────────────────────────


@asynccontextmanager
async def lifespan(app: FastAPI):
    for attempt in range(15):
        try:
            async with engine.begin() as conn:
                await conn.run_sync(Base.metadata.create_all)
            print("✅ Database ready.")
            break
        except Exception as exc:
            if attempt == 14:
                raise RuntimeError(f"DB connection failed: {exc}") from exc
            print(f"⏳ Waiting for DB... ({attempt + 1}/15)")
            await asyncio.sleep(2)
    yield
    await engine.dispose()


app = FastAPI(title="VoiceNote AI API", version="5.0.0", lifespan=lifespan)

# Configurar limitador de peticiones en la aplicación
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)

# CORS Estricto para Producción
app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PATCH", "DELETE", "OPTIONS"],
    allow_headers=["*"],
)


# ── DB Dependency ──────────────────────────────────────────────────────────────


async def get_db() -> AsyncGenerator[AsyncSession, None]:
    async with AsyncSessionLocal() as session:
        try:
            yield session
        except Exception:
            await session.rollback()
            raise


# ── Auth ───────────────────────────────────────────────────────────────────────


def hash_password(p: str) -> str:
    return bcrypt.hashpw(p.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def verify_password(plain: str, hashed: str) -> bool:
    try:
        return bcrypt.checkpw(plain.encode("utf-8"), hashed.encode("utf-8"))
    except Exception:
        return False


def create_token(user_id: int, email: str) -> str:
    exp = datetime.now(timezone.utc) + timedelta(days=TOKEN_EXPIRE_DAYS)
    return jwt.encode(
        {"sub": str(user_id), "email": email, "exp": exp},
        JWT_SECRET,
        algorithm=ALGORITHM,
    )


async def get_current_user(
    credentials: HTTPAuthorizationCredentials = Depends(http_bearer),
    db: AsyncSession = Depends(get_db),
) -> User:
    try:
        payload = jwt.decode(
            credentials.credentials, JWT_SECRET, algorithms=[ALGORITHM]
        )
        user_id = int(payload["sub"])
    except (JWTError, KeyError, ValueError):
        raise HTTPException(status_code=401, detail="Token inválido o expirado")
    result = await db.execute(select(User).where(User.id == user_id))
    user = result.scalar_one_or_none()
    if not user:
        raise HTTPException(status_code=401, detail="Usuario no encontrado")
    return user


# ── Schemas ────────────────────────────────────────────────────────────────────


class RegisterRequest(BaseModel):
    name: str
    email: str
    password: str


class LoginRequest(BaseModel):
    email: str
    password: str


class UpdateProfileRequest(BaseModel):
    name: Optional[str] = None
    current_password: Optional[str] = None
    new_password: Optional[str] = None


class UserOut(BaseModel):
    id: int
    name: str
    email: str
    created_at: datetime
    model_config = {"from_attributes": True}


class AuthResponse(BaseModel):
    access_token: str
    user: UserOut


class MeetingCreate(BaseModel):
    title: str
    transcript: str = ""
    duration: int = 0


class MeetingUpdate(BaseModel):
    title: Optional[str] = None
    transcript: Optional[str] = None
    summary: Optional[str] = None
    duration: Optional[int] = None


class MeetingOut(BaseModel):
    id: int
    title: str
    transcript: str
    summary: str
    duration: int
    share_token: Optional[str]
    created_at: datetime
    model_config = {"from_attributes": True}


class ChatRequest(BaseModel):
    query: str


# ── Health ─────────────────────────────────────────────────────────────────────


@app.get("/api/health")
async def health():
    return {
        "status": "ok",
        "version": "5.0.0",
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "groq_enabled": bool(GROQ_API_KEY),
        "claude_enabled": bool(CLAUDE_API_KEY),
        "gemini_enabled": bool(GEMINI_API_KEY),
    }


# ── Auth endpoints ─────────────────────────────────────────────────────────────


@app.post("/api/auth/register", response_model=AuthResponse, status_code=201)
@limiter.limit("5/minute")
async def register(
    request: Request, data: RegisterRequest, db: AsyncSession = Depends(get_db)
):
    if not data.name.strip():
        raise HTTPException(400, "El nombre es requerido")
    if "@" not in data.email:
        raise HTTPException(400, "Email inválido")
    if len(data.password) < 6:
        raise HTTPException(400, "La contraseña debe tener al menos 6 caracteres")
    r = await db.execute(select(User).where(User.email == data.email.lower()))
    if r.scalar_one_or_none():
        raise HTTPException(400, "Este email ya está registrado")
    user = User(
        name=data.name.strip(),
        email=data.email.lower().strip(),
        password_hash=hash_password(data.password),
    )
    db.add(user)
    await db.commit()
    await db.refresh(user)
    return AuthResponse(
        access_token=create_token(user.id, user.email),
        user=UserOut.model_validate(user),
    )


@app.post("/api/auth/login", response_model=AuthResponse)
@limiter.limit("10/minute")
async def login(
    request: Request, data: LoginRequest, db: AsyncSession = Depends(get_db)
):
    r = await db.execute(select(User).where(User.email == data.email.lower()))
    user = r.scalar_one_or_none()
    if not user or not verify_password(data.password, user.password_hash):
        raise HTTPException(401, "Email o contraseña incorrectos")
    return AuthResponse(
        access_token=create_token(user.id, user.email),
        user=UserOut.model_validate(user),
    )


@app.get("/api/auth/me", response_model=UserOut)
async def me(current_user: User = Depends(get_current_user)):
    return current_user


@app.patch("/api/auth/me", response_model=UserOut)
async def update_profile(
    data: UpdateProfileRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if data.name is not None:
        if not data.name.strip():
            raise HTTPException(400, "El nombre no puede estar vacío")
        current_user.name = data.name.strip()

    if data.new_password:
        if not data.current_password:
            raise HTTPException(400, "Debes proporcionar la contraseña actual")
        if not verify_password(data.current_password, current_user.password_hash):
            raise HTTPException(400, "Contraseña actual incorrecta")
        if len(data.new_password) < 6:
            raise HTTPException(
                400, "La nueva contraseña debe tener al menos 6 caracteres"
            )
        current_user.password_hash = hash_password(data.new_password)

    await db.commit()
    await db.refresh(current_user)
    return current_user


@app.delete("/api/auth/account", status_code=204)
async def delete_account(
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await db.delete(current_user)
    await db.commit()


# ── Meeting CRUD ───────────────────────────────────────────────────────────────


@app.get("/api/meetings", response_model=list[MeetingOut])
async def list_meetings(
    skip: int = 0,
    limit: int = 15,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Meeting)
        .where(Meeting.user_id == current_user.id)
        .order_by(Meeting.created_at.desc())
        .offset(skip)
        .limit(limit)
    )
    return r.scalars().all()


@app.post("/api/meetings", response_model=MeetingOut, status_code=201)
async def create_meeting(
    data: MeetingCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    m = Meeting(user_id=current_user.id, **data.model_dump())
    db.add(m)
    await db.commit()
    await db.refresh(m)
    return m


@app.get("/api/meetings/{meeting_id}", response_model=MeetingOut)
async def get_meeting(
    meeting_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Meeting).where(
            Meeting.id == meeting_id, Meeting.user_id == current_user.id
        )
    )
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Reunión no encontrada")
    return m


@app.patch("/api/meetings/{meeting_id}", response_model=MeetingOut)
async def update_meeting(
    meeting_id: int,
    data: MeetingUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Meeting).where(
            Meeting.id == meeting_id, Meeting.user_id == current_user.id
        )
    )
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Reunión no encontrada")
    for k, v in data.model_dump(exclude_none=True).items():
        setattr(m, k, v)
    await db.commit()
    await db.refresh(m)
    return m


@app.delete("/api/meetings/{meeting_id}", status_code=204)
async def delete_meeting(
    meeting_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Meeting).where(
            Meeting.id == meeting_id, Meeting.user_id == current_user.id
        )
    )
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Reunión no encontrada")
    await db.delete(m)
    await db.commit()


# ── Audio helpers ──────────────────────────────────────────────────────────────


def _convert_to_wav_sync(audio_bytes: bytes) -> bytes:
    import os as _os

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = _os.path.join(tmpdir, "input.audio")
        output_path = _os.path.join(tmpdir, "output.wav")
        with open(input_path, "wb") as fh:
            fh.write(audio_bytes)
        proc = subprocess.run(
            [
                "ffmpeg",
                "-y",
                "-i",
                input_path,
                "-ar",
                "16000",
                "-ac",
                "1",
                "-f",
                "wav",
                output_path,
            ],
            capture_output=True,
            timeout=120,
        )
        if proc.returncode != 0:
            print(f"[ffmpeg] warning: {proc.stderr.decode()[:200]}")
            return audio_bytes
        with open(output_path, "rb") as fh:
            wav = fh.read()
    print(f"[ffmpeg] {len(audio_bytes)}B → {len(wav)}B WAV")
    return wav


def _apply_glossary(text: str) -> str:
    for wrong, correct in GLOSSARY.items():
        text = re.sub(rf"\b{re.escape(wrong)}\b", correct, text, flags=re.IGNORECASE)
    return text


def _clean_text(text: str) -> str:
    text = FILLERS.sub(" ", text)
    text = re.sub(r"\b(\w+)(\s+\1)+\b", r"\1", text, flags=re.IGNORECASE)
    text = re.sub(r" {2,}", " ", text).strip()
    text = _apply_glossary(text)
    if text and text[-1] not in ".!?,;:":
        text += "."
    if text:
        text = text[0].upper() + text[1:]
    return text


def _segment_into_speakers(
    text: str, verbose_segments: list | None = None
) -> list[dict]:
    PAUSE_THRESHOLD = 1.2
    speaker_cycle = ["Hablante 1", "Hablante 2", "Hablante 3", "Hablante 4"]
    result: list[dict] = []

    if verbose_segments:
        speaker_idx = 0
        prev_end = 0.0
        cur_speaker = speaker_cycle[0]
        buf_text = ""
        buf_start = 0.0
        buf_end = 0.0

        for seg in verbose_segments:
            start = seg.get("start", 0.0)
            end = seg.get("end", start + 1.0)
            text = (seg.get("text") or "").strip()
            if not text:
                continue

            if buf_text and (start - prev_end) > PAUSE_THRESHOLD:
                cleaned = _clean_text(buf_text.strip())
                if cleaned and cleaned != ".":
                    result.append(
                        {
                            "speaker": cur_speaker,
                            "text": cleaned,
                            "start": buf_start,
                            "end": buf_end,
                            "timestamp": f"{int(buf_start//60):02d}:{int(buf_start%60):02d}",
                        }
                    )
                speaker_idx = (speaker_idx + 1) % len(speaker_cycle)
                cur_speaker = speaker_cycle[speaker_idx]
                buf_text = ""
                buf_start = start

            if not buf_text:
                buf_start = start
            buf_text += " " + text
            buf_end = end
            prev_end = end

        if buf_text.strip():
            cleaned = _clean_text(buf_text.strip())
            if cleaned and cleaned != ".":
                result.append(
                    {
                        "speaker": cur_speaker,
                        "text": cleaned,
                        "start": buf_start,
                        "end": buf_end,
                        "timestamp": f"{int(buf_start//60):02d}:{int(buf_start%60):02d}",
                    }
                )
        return result

    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    speaker_idx = 0

    for i, sentence in enumerate(sentences):
        sentence = sentence.strip()
        if not sentence:
            continue
        cleaned = _clean_text(sentence)
        if not cleaned or cleaned == ".":
            continue

        if i > 0 and i % 3 == 0:
            speaker_idx = (speaker_idx + 1) % len(speaker_cycle)

        result.append(
            {
                "speaker": speaker_cycle[speaker_idx],
                "text": cleaned,
                "start": float(i * 3),
                "end": float(i * 3 + 2),
                "timestamp": f"00:{(i*3):02d}",
            }
        )

    return result


def _extract_auto_title(segments: list[dict]) -> str:
    return ""


async def _transcribe_with_groq(
    audio_bytes: bytes, filename: str = "audio.wav"
) -> dict:
    groq = get_groq()
    print(f"[Groq] Transcribing {len(audio_bytes)}B ({filename})...")

    try:
        loop = asyncio.get_event_loop()

        def _call():
            return groq.audio.transcriptions.create(
                file=(filename, audio_bytes),
                model="whisper-large-v3",
                response_format="verbose_json",
                language=None,
                prompt=(
                    "Conversación en español latinoamericano con múltiples hablantes. "
                    "Distingue claramente los cambios de turno de palabra. "
                    "Transcribe con puntuación correcta y frases completas."
                ),
                temperature=0.0,
            )

        result = await loop.run_in_executor(None, _call)

        raw_text = result.text or ""
        verbose_segs = getattr(result, "segments", None) or []

        seg_dicts = []
        for s in verbose_segs:
            if hasattr(s, "start"):
                seg_dicts.append({"start": s.start, "end": s.end, "text": s.text})
            elif isinstance(s, dict):
                seg_dicts.append(s)

        print(f"[Groq] ✅ {len(seg_dicts)} raw segments, {len(raw_text)} chars")

        segments = _segment_into_speakers(raw_text, seg_dicts if seg_dicts else None)
        auto_title = _extract_auto_title(segments)

        return {
            "segments": segments,
            "text": raw_text,
            "auto_title": auto_title,
            "error": None,
        }

    except Exception as e:
        err = str(e)
        print(f"[Groq] ❌ {err[:300]}")
        return {"segments": [], "text": "", "auto_title": "", "error": err[:300]}


# ── Real-time audio chunk → Groq ──────────────────────────────────────────────


@app.post("/api/transcribe/audio/{meeting_id}")
async def transcribe_audio_chunk(
    request: Request,
    meeting_id: int,
    audio: UploadFile = File(...),
    token: str = Query(...),
    offset: float = Query(0.0),
    db: AsyncSession = Depends(get_db),
):
    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[ALGORITHM])
        user_id = int(payload["sub"])
    except (JWTError, KeyError, ValueError):
        raise HTTPException(401, "Token inválido")
    r = await db.execute(
        select(Meeting).where(Meeting.id == meeting_id, Meeting.user_id == user_id)
    )
    if not r.scalar_one_or_none():
        raise HTTPException(404, "Reunión no encontrada")

    audio_bytes = await audio.read()
    print(f"[Chunk] {len(audio_bytes)//1024}KB received | offset={offset:.1f}s")

    loop = asyncio.get_event_loop()
    wav = await loop.run_in_executor(None, _convert_to_wav_sync, audio_bytes)

    result = await _transcribe_with_groq(wav, "chunk.wav")

    TOLERANCE = 0.5
    if offset > 0 and result.get("segments"):
        result["segments"] = [
            s for s in result["segments"] if s.get("start", 0) >= (offset - TOLERANCE)
        ]
        seen_texts: set[str] = set()
        deduped = []
        for seg in result["segments"]:
            key = seg.get("text", "").strip().lower()[:60]
            if key and key not in seen_texts:
                seen_texts.add(key)
                deduped.append(seg)
        result["segments"] = deduped

    print(
        f"[Chunk] → {len(result.get('segments', []))} new segments after offset={offset:.1f}s"
    )
    return result


# ── File Upload (Background Task) ──────────────────────────────────────────────

SUPPORTED = {
    ".mp3",
    ".ogg",
    ".wav",
    ".m4a",
    ".flac",
    ".webm",
    ".mp4",
    ".mov",
    ".avi",
    ".mkv",
}


async def process_audio_background(
    meeting_id: int, file_bytes: bytes, original_name: str
):
    print(f"[BG Task] Procesando archivo {original_name}...")
    try:
        loop = asyncio.get_event_loop()
        wav = await loop.run_in_executor(None, _convert_to_wav_sync, file_bytes)
        result = await _transcribe_with_groq(wav, "upload.wav")

        async with AsyncSessionLocal() as db:
            r = await db.execute(select(Meeting).where(Meeting.id == meeting_id))
            m = r.scalar_one_or_none()
            if not m:
                return

            if result.get("error") and not result.get("segments"):
                m.transcript = f"[ERROR] {result['error']}"
            else:
                segments = result.get("segments", [])
                lines = [f"{s['speaker']}: {s['text']}" for s in segments]
                m.transcript = "\n".join(lines) if lines else "[VACIO]"
                m.duration = max(1, len(wav) // 32000)

            await db.commit()
            print(f"[BG Task] ✅ Finalizado {original_name}.")
    except Exception as e:
        print(f"[BG Task] ❌ Error catastrófico: {e}")
        async with AsyncSessionLocal() as db:
            r = await db.execute(select(Meeting).where(Meeting.id == meeting_id))
            m = r.scalar_one_or_none()
            if m:
                m.transcript = f"[ERROR] Falla interna al procesar el archivo."
                await db.commit()


@app.post("/api/transcribe/upload", response_model=MeetingOut, status_code=202)
@limiter.limit("20/hour")
async def upload_and_transcribe(
    request: Request,
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    import os as _os

    original_name = file.filename or "upload"
    ext = _os.path.splitext(original_name)[1].lower()

    if ext not in SUPPORTED:
        raise HTTPException(415, f"Formato no soportado: '{ext}'")

    file_bytes = await file.read()
    if len(file_bytes) < 500:
        raise HTTPException(400, "El archivo está vacío")

    print(
        f"[Upload Request] Recibido {original_name} | {len(file_bytes)}B. Enviando a BG task."
    )

    raw_name = _os.path.splitext(original_name)[0][:80].strip()
    clean_name = re.sub(r"[_\-]+", " ", raw_name).strip()
    title = clean_name if clean_name else "Reunión sin título"

    meeting = Meeting(
        user_id=current_user.id,
        title=title,
        transcript="[PROCESANDO]",
        summary="",
        duration=0,
    )
    db.add(meeting)
    await db.commit()
    await db.refresh(meeting)

    background_tasks.add_task(
        process_audio_background, meeting.id, file_bytes, original_name
    )

    return meeting


# ── Share ──────────────────────────────────────────────────────────────────────


@app.post("/api/meetings/{meeting_id}/share")
async def share_meeting(
    meeting_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Meeting).where(
            Meeting.id == meeting_id, Meeting.user_id == current_user.id
        )
    )
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Reunión no encontrada")
    if not m.share_token:
        m.share_token = secrets.token_urlsafe(32)
        await db.commit()
        await db.refresh(m)
    return {"share_token": m.share_token}


@app.get("/api/share/{token}", response_model=MeetingOut)
async def get_shared(token: str, db: AsyncSession = Depends(get_db)):
    r = await db.execute(select(Meeting).where(Meeting.share_token == token))
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Enlace no válido")
    return m


# ── AI Chat (Gemini) ──────────────────────────────────────────────────────────


@app.post("/api/meetings/{meeting_id}/chat")
@limiter.limit("30/minute")
async def chat_meeting(
    request: Request,
    meeting_id: int,
    data: ChatRequest,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if not GEMINI_API_KEY:
        raise HTTPException(503, "Configura GEMINI_API_KEY en el .env")

    r = await db.execute(
        select(Meeting).where(
            Meeting.id == meeting_id, Meeting.user_id == current_user.id
        )
    )
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Reunión no encontrada")

    if m.transcript == "[PROCESANDO]":
        raise HTTPException(400, "El audio aún se está procesando.")

    prompt = f"""Basándote EXCLUSIVAMENTE en la siguiente transcripción de una reunión, responde la pregunta del usuario. Si la respuesta a la pregunta no se encuentra en el texto de la reunión, responde amablemente que no tienes esa información basada en la conversación.

TRANSCRIPCIÓN:
{m.transcript}

PREGUNTA DEL USUARIO:
{data.query}
"""

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            resp = await client.post(
                f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}",
                headers={"Content-Type": "application/json"},
                json={
                    "contents": [{"parts": [{"text": prompt}]}],
                    "generationConfig": {
                        "temperature": 0.2,
                    },
                },
            )
            if resp.status_code == 200:
                resp_data = resp.json()
                answer = (
                    resp_data.get("candidates", [{}])[0]
                    .get("content", {})
                    .get("parts", [{}])[0]
                    .get("text", "")
                    .strip()
                )
                return {"answer": answer}
            else:
                raise HTTPException(502, f"Error de Gemini: {resp.status_code}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(502, f"Falla de conexión: {str(e)}")


# ── AI Summary ─────────────────────────────────────────────────────────────────

SUMMARY_PROMPT = """Eres un experto en análisis de reuniones empresariales. 
Analiza la siguiente transcripción y responde EXACTAMENTE en este formato estructurado:

RESUMEN EJECUTIVO:
[2-3 oraciones concisas que capturen la esencia de la reunión]

PUNTOS CLAVE:
• [Punto clave 1 con contexto relevante]
• [Punto clave 2 con contexto relevante]
• [Punto clave 3 con contexto relevante]

DECISIONES TOMADAS:
• [Decisión 1, o "Ninguna decisión formal identificada" si no hay]

ACCIONES A SEGUIR:
• [Acción 1 – Responsable: nombre si se menciona]
• [Acción 2 – Responsable: nombre si se menciona]

PRÓXIMOS PASOS:
[1-2 oraciones sobre qué debe suceder después de esta reunión]

TRANSCRIPCIÓN:
{transcript}"""


@app.post("/api/meetings/{meeting_id}/summarize")
@limiter.limit("30/hour")
async def summarize_meeting(
    request: Request,
    meeting_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    if not CLAUDE_API_KEY and not GEMINI_API_KEY:
        raise HTTPException(503, "Configura CLAUDE_API_KEY o GEMINI_API_KEY en el .env")

    r = await db.execute(
        select(Meeting).where(
            Meeting.id == meeting_id, Meeting.user_id == current_user.id
        )
    )
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Reunión no encontrada")

    if m.transcript == "[PROCESANDO]":
        raise HTTPException(400, "El audio aún se está procesando.")

    transcript_words = len(m.transcript.strip().split()) if m.transcript else 0
    if transcript_words < 10:
        raise HTTPException(
            400,
            f"Transcripción muy corta ({transcript_words} palabras). Graba al menos 10 palabras para generar un resumen.",
        )

    prompt = SUMMARY_PROMPT.format(transcript=m.transcript)
    summary_text = ""

    if CLAUDE_API_KEY:
        try:
            async with httpx.AsyncClient(timeout=45.0) as client:
                resp = await client.post(
                    "https://api.anthropic.com/v1/messages",
                    headers={
                        "x-api-key": CLAUDE_API_KEY,
                        "anthropic-version": "2023-06-01",
                        "content-type": "application/json",
                    },
                    json={
                        "model": "claude-haiku-4-5-20251001",
                        "max_tokens": 1500,
                        "messages": [{"role": "user", "content": prompt}],
                    },
                )
                if resp.status_code == 200:
                    data = resp.json()
                    summary_text = data["content"][0]["text"].strip()
                    print(f"[Claude] ✅ Summary generated ({len(summary_text)} chars)")
                else:
                    print(f"[Claude] ⚠️ {resp.status_code}: {resp.text[:200]}")
        except Exception as e:
            print(f"[Claude] ❌ {str(e)[:200]}")

    if not summary_text and GEMINI_API_KEY:
        try:
            async with httpx.AsyncClient(timeout=30.0) as client:
                resp = await client.post(
                    f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key={GEMINI_API_KEY}",
                    headers={"Content-Type": "application/json"},
                    json={
                        "contents": [{"parts": [{"text": prompt}]}],
                        "generationConfig": {
                            "maxOutputTokens": 1500,
                            "temperature": 0.3,
                        },
                    },
                )
                if resp.status_code == 200:
                    data = resp.json()
                    summary_text = (
                        data.get("candidates", [{}])[0]
                        .get("content", {})
                        .get("parts", [{}])[0]
                        .get("text", "")
                        .strip()
                    )
                    print(f"[Gemini] ✅ Summary generated ({len(summary_text)} chars)")
                elif resp.status_code == 403:
                    raise HTTPException(403, "GEMINI_API_KEY inválida")
                else:
                    print(f"[Gemini] ⚠️ {resp.status_code}: {resp.text[:200]}")
        except HTTPException:
            raise
        except Exception as e:
            print(f"[Gemini] ❌ {str(e)[:200]}")

    if not summary_text:
        raise HTTPException(
            502, "No se pudo generar el resumen. Verifica las API keys en el .env"
        )

    m.summary = summary_text
    await db.commit()
    await db.refresh(m)
    return {"summary": summary_text}


# ── Export (PDF/DOCX) ──────────────────────────────────────────────────────────


def create_pdf(meeting: Meeting) -> bytes:
    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=24,
        textColor=colors.HexColor("#1e1b4b"),
        spaceAfter=10,
    )
    meta_style = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=10,
        textColor=colors.HexColor("#64748b"),
        spaceAfter=20,
    )
    h2_style = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=16,
        textColor=colors.HexColor("#4338ca"),
        spaceBefore=20,
        spaceAfter=10,
    )
    spk_style = ParagraphStyle(
        "Speaker",
        parent=styles["Normal"],
        fontSize=10,
        fontName="Helvetica-Bold",
        textColor=colors.HexColor("#4f46e5"),
        spaceBefore=10,
        spaceAfter=2,
    )
    text_style = ParagraphStyle(
        "Text",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#334155"),
        leading=16,
    )
    summary_style = ParagraphStyle(
        "Summary",
        parent=styles["Normal"],
        fontSize=11,
        textColor=colors.HexColor("#334155"),
        leading=16,
        leftIndent=10,
        rightIndent=10,
    )

    elements: list[Flowable] = []

    elements.append(Paragraph(meeting.title, title_style))
    date_str = meeting.created_at.strftime("%Y-%m-%d %H:%M")
    mins = meeting.duration // 60
    elements.append(
        Paragraph(f"Fecha: {date_str} | Duración: {mins} minutos", meta_style)
    )
    elements.append(
        HRFlowable(
            width="100%", thickness=1, color=colors.HexColor("#e2e8f0"), spaceAfter=20
        )
    )

    if meeting.summary:
        elements.append(Paragraph("Resumen Ejecutivo", h2_style))
        for line in meeting.summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.isupper() and len(line) < 50:
                s = ParagraphStyle(
                    "SH", parent=spk_style, textColor=colors.HexColor("#7c3aed")
                )
                elements.append(Paragraph(line, s))
            else:
                elements.append(Paragraph(line, summary_style))
        elements.append(Spacer(1, 20))
        elements.append(
            HRFlowable(
                width="100%",
                thickness=1,
                color=colors.HexColor("#e2e8f0"),
                spaceAfter=20,
            )
        )

    elements.append(Paragraph("Transcripción", h2_style))
    lines = [L for L in meeting.transcript.split("\n") if L.strip()]

    for line in lines:
        if ":" in line:
            spk, txt = line.split(":", 1)
            elements.append(
                KeepTogether(
                    [
                        Paragraph(spk.strip(), spk_style),
                        Paragraph(txt.strip(), text_style),
                    ]
                )
            )
        else:
            elements.append(Paragraph(line.strip(), text_style))

    doc.build(elements)
    return buffer.getvalue()


def create_docx(meeting: Meeting) -> bytes:
    doc = Document()

    title = doc.add_heading(meeting.title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_str = meeting.created_at.strftime("%Y-%m-%d %H:%M")
    mins = meeting.duration // 60
    meta = doc.add_paragraph(f"Fecha: {date_str} | Duración: {mins} minutos")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta.runs:
        run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    if meeting.summary:
        h2 = doc.add_heading("Resumen Ejecutivo", level=1)
        for run in h2.runs:
            run.font.color.rgb = RGBColor(124, 58, 237)

        for line in meeting.summary.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            if line.isupper() and len(line) < 50:
                r = p.add_run(line)
                r.bold = True
                r.font.color.rgb = RGBColor(124, 58, 237)
            else:
                p.add_run(line)

        doc.add_page_break()

    h2 = doc.add_heading("Transcripción", level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(79, 70, 229)

    lines = [L for L in meeting.transcript.split("\n") if L.strip()]
    for line in lines:
        if ":" in line:
            spk, txt = line.split(":", 1)
            p = doc.add_paragraph()
            r_spk = p.add_run(spk.strip() + ":")
            r_spk.bold = True
            r_spk.font.color.rgb = RGBColor(79, 70, 229)
            p.add_run(" " + txt.strip())
        else:
            doc.add_paragraph(line.strip())

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


import urllib.parse


@app.get("/api/meetings/{meeting_id}/export/{format}")
async def export_meeting(
    meeting_id: int,
    format: str,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    r = await db.execute(
        select(Meeting).where(
            Meeting.id == meeting_id, Meeting.user_id == current_user.id
        )
    )
    m = r.scalar_one_or_none()
    if not m:
        raise HTTPException(404, "Reunión no encontrada")

    # Limpiamos el nombre del archivo para evitar errores HTTP en cabeceras
    safe_title = re.sub(r"[^a-zA-Z0-9_\-]", "", m.title.replace(" ", "_"))
    if not safe_title:
        safe_title = f"Reunion_{m.id}"

    encoded_filename = urllib.parse.quote(f"{safe_title}.{format}")

    if format == "pdf":
        pdf_bytes = create_pdf(m)
        return StreamingResponse(
            BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            },
        )
    elif format == "docx":
        docx_bytes = create_docx(m)
        return StreamingResponse(
            BytesIO(docx_bytes),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"
            },
        )
    else:
        raise HTTPException(400, "Formato no soportado")


# ── Donaciones (Mercado Pago) ──────────────────────────────────────────────────
import mercadopago
import os

# Usamos .strip() para limpiar espacios invisibles del .env
MP_ACCESS_TOKEN = os.getenv("MP_ACCESS_TOKEN", "").strip()
mp_sdk = mercadopago.SDK(MP_ACCESS_TOKEN) if MP_ACCESS_TOKEN else None


class DonationRequest(BaseModel):
    amount: int


@app.post("/api/donate")
@limiter.limit("5/minute")
async def create_donation_preference(request: Request, data: DonationRequest):
    if not mp_sdk:
        raise HTTPException(
            503, "Mercado Pago no configurado en .env (falta MP_ACCESS_TOKEN)"
        )

    preference_data = {
        "items": [
            {
                "id": "donacion",
                "title": "Aporte a VoiceNote AI",
                "quantity": 1,
                "unit_price": float(data.amount),
            }
        ]
    }

    try:
        preference_response = mp_sdk.preference().create(preference_data)
        pref = preference_response.get("response", {})

        if "id" in pref:
            return {"preference_id": pref["id"]}
        else:
            # AQUÍ ESTÁ LA CLAVE: Si falla, mostramos el error real de MP
            error_msg = pref.get("message", "Error desconocido de Mercado Pago")
            raise HTTPException(500, f"Detalle MP: {error_msg}")

    except Exception as e:
        if isinstance(e, HTTPException):
            raise e
        raise HTTPException(500, f"Error interno: {str(e)}")
